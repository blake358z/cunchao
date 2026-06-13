from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT_MD = "村超APP小程序立项研究.md"
OUT_DOCX = "村超APP小程序立项研究.docx"


sources = [
    ("村超官网：赛事情况", "https://gzcunchao.com/ssqk"),
    ("贵州省政府专题：贵州村超来了-联赛赛程", "https://www.guizhou.gov.cn/ztzl/gzcfa/lssc/"),
    ("人民日报：2026村超玩法更多，走向立体丰盈", "https://paper.people.com.cn/rmrb/pc/content/202604/10/content_30150336.html"),
    ("人民网：2026年第四届村超在贵州榕江开赛", "https://ent.people.com.cn/n1/2026/0103/c1012-40637810.html"),
    ("新华网：首届全国赛落幕 村超聚力再出发", "https://www.news.cn/sports/20250824/9a2de09aca3f490fb410d82e1dc0f52e/c.html"),
    ("新华网：榕江推出足球发展十年规划", "https://www.news.cn/sports/20250531/671e0390e4fa4165a71b579a0de5007a/c.html"),
    ("懂球帝 Google Play 应用页", "https://play.google.com/store/apps/details?hl=zh&id=com.soccer.funs.news"),
    ("懂球帝官网", "https://www.dongqiudi.com/"),
    ("虎扑产品体验分析报告", "https://www.woshipm.com/evaluating/1605762.html"),
    ("虎扑 APP 产品分析报告", "https://www.woshipm.com/evaluating/3757410.html"),
    ("体育多 Google Play 应用页", "https://play.google.com/store/apps/details?hl=zh_TW&id=com.tiyuduo.app"),
]


sections = [
    ("一、研究结论摘要", [
        "村超的产品机会不只是“看比赛”，而是把乡村足球赛事、地方文化、球迷社区、短视频传播、线下文旅消费连接起来。APP/小程序应同时服务游客、球迷、球队/球员、赛事运营者和本地商户。",
        "赛制层面，榕江本地村超已经从2023年的20支队伍扩容到2026年的137支村级球队；2026年本地联赛包含预选赛、40强、20强、8强和淘汰赛决赛等阶段。省级冠军赛和全国赛则进一步把“村超”扩展为跨区域赛事矩阵。",
        "产品层面，虎扑可借鉴的是论坛文化、赛事热线、评分/亮帖机制和球队/球员数据榜；懂球帝可借鉴的是足球垂直资讯、赛程比分、积分/射手榜、赛前赛后数据与球迷圈子。村超产品需要在此基础上加入“村味”：村寨身份、民族文化展演、农特产、旅行攻略、线下预约和志愿服务。",
    ]),
    ("二、村超赛事与赛制梳理", [
        "基础定位：村超正式名称常见为榕江（三宝侗寨）和美乡村足球超级联赛，发源于贵州省黔东南州榕江县，以村民组织、村级球队、非职业球员和乡土奖品为显著特征。",
        "2026榕江本地联赛：公开报道显示，第四届村超有137支村级球队、2694名非职业球员参赛，覆盖榕江县20个乡镇（街道）的村寨与社区。预选赛共422场，分20个小组，每组6至7支队伍，采用小组循环积分制，每组前2名晋级40强。40强后进行交叉单场淘汰，产生20强；20强分2组循环，各组前4进入8强，再单场淘汰决出年度冠军。",
        "2026时间线：1月至4月为年度联赛预选赛；4月底开启20强赛；预计6月至7月进行总决赛阶段，约98场对决决出冠军。所有场次延续零门票传统，日常可直接入场，高峰期可能采用扫码预约与限流。",
        "贵州省级冠军赛：2026年贵州村超冠军赛由9个市州代表队参加，5月16日至8月22日进行，共30场，采用主客场双循环并分省级分组赛、附加赛、总决赛排位赛、总决赛四个阶段。",
        "全国赛：首届全国赛在2025年形成“各赛区选拔+榕江总决赛”的模式。2025年3月至7月为全国各赛区赛事，7月至8月在榕江举行总决赛；新华社报道其覆盖34个省级行政区、36个赛区，从683支球队中选拔51支球队赴榕江参赛。",
        "长期方向：榕江2025-2035足球发展方案提出“村超”社会足球、“班超”校园足球、“逐梦”职业足球三线体系，并规划本地联赛、全国赛事、国际交流赛事三级赛事矩阵。",
    ]),
    ("三、目标用户与核心需求", [
        "本地球迷/村民：查看赛程、支持本村球队、参与助威、看直播回放、分享短视频、参与投票和讨论。",
        "外地游客：查比赛日程、预约入场、交通住宿、吃喝玩乐、非遗/演出时间、周边市集和农特产购买。",
        "球队/球员：球队主页、球员档案、报名与资格资料、赛程提醒、战绩数据、照片视频素材沉淀。",
        "赛事运营者：排赛、比分录入、红黄牌/进球/换人记录、积分榜自动计算、内容发布、志愿者和现场流量管理。",
        "本地商户/文旅机构：赛事流量转化、活动曝光、优惠券、路线推荐、农特产电商或团购。",
    ]),
    ("四、竞品功能拆解", [
        "虎扑：核心是“体育资讯+赛事+社区”。可借鉴底部主导航、赛事页、球队/球员榜单、赛前赛中赛后直播间、评论亮帖/点灭/举报、话题圈和社区答题/规则教育机制。对村超而言，虎扑的社区氛围治理和赛后讨论留存尤其值得参考。",
        "懂球帝：核心是“足球垂直内容+赛事数据”。公开应用页强调资讯、比赛、聊球、数据；官网有赛程、积分、射手榜、历史等模块。可借鉴足球专业信息架构：比赛详情、赛前分析、赛况数据、球员评分、积分榜、射手榜、助攻榜、赛程实时更新。",
        "体育多/即时比分类产品：强调即时比分、进球/红黄牌/换人事件、比赛提醒、技术统计、历史对战和数据分析。村超第一阶段可采用轻量数据模型，不必一次上高阶技术统计。",
        "村超官网/官方专题：已有赛事情况、新闻、精彩视频、预约、旅游攻略、住宿交通等信息，说明官方信息入口已经具备雏形。新产品需要做的是移动端体验、数据结构化、社区互动和文旅转化闭环。",
    ]),
    ("五、建议的信息架构", [
        "首页：今日赛事、热门直播/回放、最新战报、热门讨论、本周文旅活动。",
        "赛事：赛程日历、分组赛/淘汰赛、比分直播、比赛详情、积分榜、射手榜、球队榜、数据榜。",
        "球队：村寨球队主页、队徽/村寨故事、球员名单、战绩、赛程、照片视频、粉丝支持。",
        "社区：赛事热线、球队圈子、村寨圈子、热门话题、亮帖、投票、评分、举报和版规。",
        "视频/新闻：战报、短视频、直播回放、人物故事、非遗展演、村寨美食、官方公告。",
        "文旅服务：预约入场、交通、住宿、地图、停车、志愿服务、周边游线路、商户优惠、农特产。",
        "我的：关注球队/球员、比赛提醒、收藏、发帖记录、订单/预约、个人认证。",
    ]),
    ("六、MVP 范围建议", [
        "MVP 1：先做微信小程序，聚焦轻量访问和线下场景。必做：赛程、比分、积分榜、球队/球员页、新闻视频、预约/交通、关注提醒。",
        "MVP 2：加入社区能力。必做：赛事热线、评论、亮帖、话题圈、举报、内容审核后台。",
        "MVP 3：运营后台。必做：球队/球员/赛程管理、比分事件录入、积分榜自动计算、内容发布、预约限流配置。",
        "MVP 4：APP 扩展。适合承载更完整的社区、短视频、会员/积分体系、商户与电商能力。",
    ]),
    ("七、数据模型初稿", [
        "赛事 Season：名称、年份、类型（本地/省级/全国/国际友谊）、阶段、状态。",
        "比赛 Match：主队、客队、时间、场地、阶段、小组、比分、状态、直播地址、回放地址。",
        "球队 Team：村寨/城市、队徽、简介、所属赛区、球员名单、战绩统计。",
        "球员 Player：姓名、号码、位置、年龄段、职业/身份标签、进球、助攻、黄牌、红牌、出场。",
        "事件 MatchEvent：进球、点球、乌龙、黄牌、红牌、换人、伤停补时、文字直播节点。",
        "内容 Content：新闻、视频、公告、攻略、人物故事、标签、关联球队/比赛。",
        "社区 Post/Comment：话题、圈子、亮帖数、举报状态、审核状态。",
        "文旅 Merchant/Route/Booking：商户、路线、活动、预约、优惠、库存/限流。",
    ]),
    ("八、下一步研究任务", [
        "用户访谈：至少覆盖本地球迷、外地游客、球队组织者、赛事运营方、本地商户五类用户。",
        "竞品截图标注：基于虎扑UI参考文档继续整理首页、赛事页、社区页、球队页、数据榜的模块级标注。",
        "PRD：把MVP拆成小程序端、APP端、运营后台、数据后台四份需求清单。",
        "原型：先做微信小程序5个核心页面：首页、赛事、比赛详情、球队详情、文旅服务。",
        "技术预研：确认数据录入方式、直播/回放来源、地图与预约能力、内容审核与合规边界。",
    ]),
]


tables = {
    "赛制矩阵": [
        ["层级", "覆盖范围", "典型赛制", "产品重点"],
        ["榕江本地联赛", "榕江县村寨/社区", "小组循环、40强、20强、8强、淘汰赛", "赛程、比分、积分、球队/球员档案"],
        ["贵州冠军赛", "贵州9个市州", "主客场双循环，分组赛、附加赛、排位赛、总决赛", "跨市州赛程、主客场、代表队专区"],
        ["全国赛", "全国赛区", "各赛区选拔，榕江总决赛", "赛区地图、晋级路径、全国球队故事"],
        ["国际/友谊交流", "国际民间球队/嘉宾", "友谊赛、交流赛", "传播、视频、文旅接待、活动专题"],
    ],
    "竞品借鉴": [
        ["产品", "强项", "可借鉴功能", "村超化改造"],
        ["虎扑", "社区文化和赛后讨论", "赛事热线、亮帖、球队/球员榜、话题圈、举报/规则机制", "村寨圈、球队助威、文明观赛答题"],
        ["懂球帝", "足球垂直资讯和赛事数据", "赛程、积分榜、射手榜、赛前/赛后数据、球员评分", "本地联赛数据、非职业球员故事、村寨榜单"],
        ["即时比分类", "实时赛况和提醒", "比分、进球、红黄牌、换人、比赛提醒", "轻量事件录入，适配现场志愿者操作"],
        ["官方村超入口", "权威信息和文旅服务", "新闻、视频、预约、交通住宿攻略", "结构化数据、移动端闭环、社区互动"],
    ],
}


def markdown():
    lines = [
        "# 村超 APP 与微信小程序立项研究",
        "",
        "生成日期：2026-06-13",
        "",
        "> 说明：用户提供的《虎扑UI参考.docx》以截图/视觉参考为主，未能抽取到正文文本；本报告将其作为后续 UI 标注参考，功能研究主要基于公开资料与竞品公开信息。",
        "",
    ]
    for title, paras in sections:
        lines.append(f"## {title}")
        for p in paras:
            lines.append(f"- {p}")
        lines.append("")
        if title == "二、村超赛事与赛制梳理":
            lines.append("### 赛制矩阵")
            for row in tables["赛制矩阵"]:
                lines.append("| " + " | ".join(row) + " |")
                if row == tables["赛制矩阵"][0]:
                    lines.append("|---|---|---|---|")
            lines.append("")
        if title == "四、竞品功能拆解":
            lines.append("### 竞品借鉴表")
            for row in tables["竞品借鉴"]:
                lines.append("| " + " | ".join(row) + " |")
                if row == tables["竞品借鉴"][0]:
                    lines.append("|---|---|---|---|")
            lines.append("")
    lines.append("## 参考来源")
    for name, url in sources:
        lines.append(f"- [{name}]({url})")
    lines.append("")
    return "\n".join(lines)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(9)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, title, data):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [Inches(1.05), Inches(1.3), Inches(2.0), Inches(2.1)]
    for i, row in enumerate(data):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            set_cell_text(cell, text, bold=(i == 0))
            cell.width = widths[j]
            if i == 0:
                set_cell_shading(cell, "E8EEF5")
    doc.add_paragraph()


def add_bullets(doc, paras):
    for p_text in paras:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(p_text)
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10.5)


def docx():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("村超 APP 与微信小程序立项研究")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 0, 0)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    r = sub.add_run("生成日期：2026-06-13｜阶段：立项研究 / 需求探索")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(85, 85, 85)

    note = doc.add_paragraph()
    note.paragraph_format.left_indent = Inches(0.15)
    note.paragraph_format.space_after = Pt(12)
    nr = note.add_run("说明：用户提供的《虎扑UI参考.docx》以截图/视觉参考为主，未能抽取到正文文本；本报告将其作为后续 UI 标注参考，功能研究主要基于公开资料与竞品公开信息。")
    nr.font.name = "Arial"
    nr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    nr.font.size = Pt(9.5)
    nr.italic = True

    for title_text, paras in sections:
        doc.add_heading(title_text, level=1)
        add_bullets(doc, paras)
        if title_text == "二、村超赛事与赛制梳理":
            add_table(doc, "赛制矩阵", tables["赛制矩阵"])
        if title_text == "四、竞品功能拆解":
            add_table(doc, "竞品借鉴表", tables["竞品借鉴"])

    doc.add_heading("参考来源", level=1)
    for name, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{name}：{url}")
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(9)

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    with open(OUT_MD, "w", encoding="utf-8") as f:
        f.write(markdown())
    docx()
    print(OUT_MD)
    print(OUT_DOCX)
