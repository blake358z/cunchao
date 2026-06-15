from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "村超MiniApp产品PRD.docx"


def style_document(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(159, 18, 57)),
        ("Heading 2", 13, RGBColor(31, 41, 55)),
        ("Heading 3", 11.5, RGBColor(55, 65, 81)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.bold = bold


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def update_version_line(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("版本："):
            paragraph.text = "版本：v0.2｜日期：2026-06-16｜阶段：MVP 原型开发 / 真实登录接入"
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10.5)
            return


def remove_existing_v02_section(doc: Document) -> None:
    start_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.startswith("16. 2026-06-16 版本进展与功能补充"):
            start_idx = idx
            break
    if start_idx is None:
        return
    for paragraph in doc.paragraphs[start_idx:]:
        p = paragraph._element
        p.getparent().remove(p)


def main() -> None:
    doc = Document(DOCX_PATH)
    style_document(doc)
    update_version_line(doc)
    remove_existing_v02_section(doc)

    doc.add_page_break()
    doc.add_heading("16. 2026-06-16 版本进展与功能补充（H5 / API / 账号体系）", level=1)
    doc.add_paragraph(
        "本节记录自上次 PRD 与 UI 设计稿同步后，产品从“设计原型”进入“可点击 H5 + API 预留 + 自动数据更新 + APK 包装”的新增范围。"
        "后续产品、前后端和运营文档以本节作为 v0.2 增量依据。"
    )

    doc.add_heading("16.1 当前可用版本与部署状态", level=2)
    add_table(
        doc,
        ["项目", "当前状态", "说明"],
        [
            ["H5 测试版", "已部署", "Vercel 测试地址为 https://cunchao-static.vercel.app，用于日常体验和合作方评审。"],
            ["前端技术栈", "已确定", "React + Vite，按数据层、业务状态、页面组合、组件样式分层，便于迁移 App / 小程序。"],
            ["Android APK", "已打包过测试包", "当前 APK 以 WebView/Capacitor 包装线上 H5，因此内容会随线上 H5 和静态数据更新。"],
            ["后端 API", "已搭建基础服务", "Fastify API 已预留赛事、球队、内容、帖子、预约和账号接口。"],
            ["GitHub 仓库", "已接入", "代码、数据更新脚本、自动化工作流均进入仓库管理。"],
            ["国内访问方案", "已有文档", "Vercel 在中国大陆访问不稳定，正式测试建议国内 OSS/COS + CDN 或香港节点过渡。"],
        ],
    )

    doc.add_heading("16.2 H5 首版体验优化", level=2)
    add_table(
        doc,
        ["模块", "已完成优化", "后续要求"],
        [
            ["首页", "焦点比赛、今天看什么、未来赛程、社区热帖、最新资讯。", "后续按关注联赛/球队做个性化排序。"],
            ["赛事", "联赛说明、今日/明日/14天筛选，比赛卡进入详情。", "结构化赛程仍需官方数据或人工维护源。"],
            ["比赛详情", "比赛事实、球队摘要、预约入场、关注比赛、参与讨论。", "生产版接入实时比分、事件流、阵容与数据员后台。"],
            ["资讯详情", "展示图片、来源、时间、正文摘要/正文、互动区。", "持续保留原文链接与来源版权标识。"],
            ["社区", "30 个模拟话题、约 1500 条讨论/回复、点赞评论收藏数据。", "上线前接入审核、举报、敏感词和运营精选。"],
            ["文旅", "按未来 14 天比赛匹配村寨/球队卡片。", "已抓取村寨资料缓存复用，缺失时再补充。"],
            ["我的", "记录本地互动、关注、评论、点赞、收藏、浏览等行为。", "登录后从服务端同步跨设备行为。"],
            ["顶部沉浸", "H5 去掉多余白色顶部状态栏，APK 尽量全屏。", "iOS/Android 真机分别做安全区适配。"],
        ],
    )

    doc.add_heading("16.3 每日数据更新工作流", level=2)
    add_table(
        doc,
        ["数据类型", "当前方案", "更新频率", "产品表现"],
        [
            ["新闻资讯", "从官方、政府专题、权威媒体等公开来源抓取，生成静态 bootstrap.json。", "每天北京时间 08:30", "新资讯排在前面，旧资讯长期保留。"],
            ["资讯图片/正文", "详情页尽量保存图片、正文摘要/正文与来源信息。", "随每日任务更新", "用户可在站内完成基本阅读，再跳转原文核验。"],
            ["未来赛程", "合并公开赛程文章与 data/manual/matches.json 的结构化比赛。", "与资讯每日任务一起更新", "赛事页展示未来 14 天对战球队、时间、地点，未开赛不显示比分。"],
            ["文旅卡片", "根据未来 14 天比赛匹配村寨、球队、场地和攻略。", "每次赛事数据更新后检查", "村寨资料已入库则复用，缺失时再补充。"],
            ["村寨图片", "为每个村寨查找更相关的村寨/目的地图片。", "文旅资料新增时执行", "避免使用与村寨不相关的随机图片。"],
            ["社区头像", "约 50 张随机头像池，模拟用户评论时随机分配。", "新增模拟内容时执行", "保留评论数量，同时提高讨论真实感。"],
        ],
    )
    doc.add_paragraph(
        "自动化发布链路：每日任务生成 apps/web/public/data/bootstrap.json，构建并验证前端；若数据有变化，自动提交 Update public data，"
        "再部署到 Vercel Production。用户打开 H5/APK 时会读取最新静态数据。"
    )

    doc.add_heading("16.4 账号体系与真实登录", level=2)
    add_table(
        doc,
        ["能力", "当前接口/实现", "说明"],
        [
            ["手机验证码发送", "POST /api/auth/phone/code", "使用腾讯云短信 SendSms，不再使用固定验证码或本地伪登录。"],
            ["手机验证码登录", "POST /api/auth/phone/login", "校验验证码后返回 AuthSession。"],
            ["微信登录", "POST /api/auth/wechat/login", "小程序环境通过 wx.login 获取 code，后端调用微信 jscode2session 换取 openid。"],
            ["当前用户", "GET /api/auth/me", "通过 Authorization: Bearer <token> 获取当前登录用户。"],
            ["退出登录", "POST /api/auth/logout", "清理当前 session。"],
            ["前端触发点", "评论、回复、收藏、关注、预约入场", "未登录时拉起登录弹层。"],
        ],
    )
    add_table(
        doc,
        ["类型", "环境变量", "用途"],
        [
            ["微信登录", "WECHAT_APP_ID, WECHAT_APP_SECRET", "调用微信 jscode2session。"],
            ["腾讯云短信", "TENCENT_SECRET_ID, TENCENT_SECRET_KEY, TENCENT_SMS_APP_ID, TENCENT_SMS_SIGN_NAME, TENCENT_SMS_TEMPLATE_ID, TENCENT_SMS_REGION", "发送真实短信验证码。"],
            ["验证码安全", "AUTH_CODE_PEPPER", "验证码哈希加盐。"],
            ["前端 API", "VITE_API_BASE_URL", "H5 指向后端 API 域名。"],
        ],
    )
    doc.add_paragraph(
        "生产化要求：当前用户、验证码和 session 仍是 Node 内存存储，适合联调和单实例测试；正式上线前必须改为数据库 + Redis，并补齐手机号/IP/设备限流、账号合并、refresh token、审计日志和隐私合规弹窗。"
    )

    doc.add_heading("16.5 后端与数据接口补充", level=2)
    add_table(
        doc,
        ["接口", "用途", "生产化备注"],
        [
            ["GET /api/bootstrap", "首屏聚合数据", "后续可由静态 JSON 过渡到 CDN 缓存 API。"],
            ["GET /api/leagues", "联赛列表", "需支持多联赛排序、地区、热度和关注态。"],
            ["GET /api/matches", "比赛列表", "需支持联赛、日期、状态、球队、场地筛选。"],
            ["GET /api/matches/:id", "比赛详情", "需接入比分、事件、阵容、评论与预约。"],
            ["GET /api/teams/:id", "球队详情", "需接入球队资料、球员、赛程、讨论和文旅。"],
            ["GET /api/content/:id", "资讯详情", "需保留来源、图片、正文、版权和互动数据。"],
            ["GET /api/posts/:id", "帖子详情", "需接入评论分页、点赞、举报、审核。"],
            ["GET /api/bookings", "可预约场次", "需和官方/第三方预约平台打通或跳转。"],
        ],
    )

    doc.add_heading("16.6 新增产品需求", level=2)
    add_bullets(
        doc,
        [
            "数据后台必须支持“自动抓取 + 人工审核 + 官方数据导入”三种数据来源，并记录每条数据的来源、更新时间、可信等级和最后审核人。",
            "资讯详情必须优先站内可读，同时保留原文链接；若版权不允许全文抓取，则至少展示标题、封面图、摘要、来源、发布时间和跳转入口。",
            "赛事页默认展示未来 14 天比赛；未开始比赛不显示比分，只显示时间、球队、场地、联赛和预约/提醒入口。",
            "文旅页主逻辑从静态攻略改为“由未来比赛驱动的村寨/球队攻略”，包括村寨介绍、球队介绍、队员介绍、旅游、住宿、饮食、交通和比赛日建议。",
            "社区模拟内容只用于测试氛围和样式，正式上线必须区分真实用户内容、运营初始化内容和机器生成种子内容，避免误导用户。",
            "用户在资讯、帖子、评论、球队、比赛、文旅攻略中的点赞、评论、收藏、关注、预约，都需要沉淀到“我的”。",
            "Android APK 当前可跟随 H5/Vercel 内容变化；若后续要上架应用市场，需要补充原生启动页、隐私政策、权限说明、推送、版本更新和备案合规。",
        ],
    )

    doc.add_heading("16.7 下一步优先级", level=2)
    add_table(
        doc,
        ["优先级", "工作项", "目标"],
        [
            ["P0", "数据持久化", "用 PostgreSQL/MySQL + Redis 替代内存用户/session/验证码，并保存互动行为。"],
            ["P0", "真实 API 部署", "为 H5、小程序和 APK 提供稳定后端域名，配置 VITE_API_BASE_URL。"],
            ["P0", "数据合作清单", "明确村超、苏超、青超等联赛的赛程、比分、球队、球员和预约数据责任方。"],
            ["P1", "小程序适配", "将微信登录、订阅消息、预约跳转、分享卡片按小程序规范落地。"],
            ["P1", "内容审核后台", "支持资讯入库审核、社区帖子/评论审核、举报处理和敏感词。"],
            ["P1", "国内部署", "准备备案域名、国内 OSS/COS + CDN，降低大陆用户访问门槛。"],
            ["P2", "数据看板", "追踪资讯更新量、比赛覆盖率、用户互动、热门球队和文旅点击。"],
        ],
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
