from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "村超MiniApp产品PRD.docx"


def style_document(doc: Document) -> None:
    styles = doc.styles
    for name, size, color in [
        ("Heading 1", 16, RGBColor(159, 18, 57)),
        ("Heading 2", 13, RGBColor(31, 41, 55)),
        ("Heading 3", 11.5, RGBColor(55, 65, 81)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def remove_existing_section(doc: Document, heading: str) -> None:
    paragraphs = list(doc.paragraphs)
    start = None
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.startswith(heading):
            start = idx
            break
    if start is None:
        return
    for paragraph in paragraphs[start:]:
        element = paragraph._element
        element.getparent().remove(element)


def set_cell(cell, value: str, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(value)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.bold = bold


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell(tbl.rows[0].cells[idx], header, True)
    for row in rows:
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value)
    doc.add_paragraph()


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document(DOCX_PATH)
    style_document(doc)
    remove_existing_section(doc, "17. 2026-06-16 管理后台与后端能力补充")

    doc.add_page_break()
    doc.add_heading("17. 2026-06-16 管理后台与后端能力补充", level=1)
    doc.add_paragraph(
        "本轮新增运营管理后台，用于支撑资讯/赛事类产品的日常运营、数据审核和增长看板。"
        "后台首版已接入 H5 /admin 路径，并新增管理后端接口和 Vercel serverless 接口。详细后端 PRD 见 docs/admin-backend-prd.md。"
    )

    doc.add_heading("17.1 后台模块", level=2)
    table(
        doc,
        ["模块", "首版能力", "后续生产化"],
        [
            ["总览", "DAU、平均观看时长、资讯 CTR、预约转化、内容池规模、模块健康度。", "接真实埋点仓库、告警和运营日报。"],
            ["内容上新", "新建资讯/视频/战报/攻略/公告草稿，查看审核队列。", "CMS 富文本、图片上传、定时发布、版本回滚。"],
            ["赛事管理", "未来比赛核验、赛程数据结构、比分事件规划。", "官方数据导入、比分录入、积分榜计算。"],
            ["球队球员", "球队资料、球员字段、球队管理员认领规划。", "授权资料库、纠错流程、敏感信息保护。"],
            ["文旅攻略", "按未来比赛匹配村寨攻略、住宿餐饮、图片素材。", "商户后台、路线编辑、攻略复用和版权管理。"],
            ["社区审核", "帖子评论巡检、治理能力清单。", "敏感词、举报工单、禁言、删除、置顶和精选。"],
            ["数据看板", "页面 PV、平均时长、CTR、预约漏斗、事件字典。", "PostHog/ClickHouse/数据仓库。"],
            ["系统设置", "角色权限、数据源、发布流、审计日志规划。", "RBAC、API 密钥、灰度发布。"],
        ],
    )

    doc.add_heading("17.2 内容生产板块", level=2)
    table(
        doc,
        ["内容类型", "前台分发位置", "关键字段"],
        [
            ["快讯/新闻", "首页信息流、资讯详情", "标题、摘要、正文、封面图、来源、发布时间、原文链接。"],
            ["赛前前瞻", "首页、赛事页、比赛详情", "对阵、看点、双方近况、预约说明。"],
            ["赛后战报", "首页、比赛详情、球队页", "比分、关键事件、最佳球员、评论热度。"],
            ["视频/图集", "首页、资讯详情、球队页", "封面、视频/图集地址、版权说明。"],
            ["球队故事", "球队页、首页专题", "村寨故事、球队历史、荣誉、人物。"],
            ["文旅攻略", "文旅页、比赛详情", "村寨介绍、路线、住宿、餐饮、交通、图片来源。"],
            ["官方公告", "首页、赛事页", "正文、有效期、发布单位。"],
            ["社区精选", "社区、首页热帖", "帖子、作者、精选理由、审核状态。"],
        ],
    )

    doc.add_heading("17.3 后端接口与埋点", level=2)
    table(
        doc,
        ["接口/事件", "类型", "用途", "当前状态"],
        [
            ["/api/admin/overview", "GET", "后台总览、模块健康度、漏斗、事件字典。", "Fastify 与 Vercel serverless 已开发。"],
            ["/api/admin/content/drafts", "POST", "创建内容草稿并进入审核队列。", "Fastify 已开发，前端首版用 localStorage 草稿。"],
            ["/api/analytics/events", "POST", "接收页面浏览、点击、预约、评论、收藏等事件。", "Fastify 与 Vercel serverless 已开发。"],
            ["page_view", "埋点", "DAU、路径分析、页面 PV。", "事件字典已定义。"],
            ["content_click", "埋点", "内容 CTR、推荐效果。", "事件字典已定义。"],
            ["match_click", "埋点", "赛事详情转化、热门比赛。", "事件字典已定义。"],
            ["booking_click", "埋点", "预约漏斗。", "事件字典已定义。"],
            ["comment_submit", "埋点", "社区活跃。", "事件字典已定义。"],
            ["favorite_toggle", "埋点", "我的沉淀、内容收藏率。", "事件字典已定义。"],
        ],
    )

    doc.add_heading("17.4 当前实现边界", level=2)
    bullets(
        doc,
        [
            "/admin 已可作为桌面管理后台入口，用于产品评审和运营流程确认。",
            "后台首版基于现有静态数据推导看板，内容草稿保存在浏览器 localStorage。",
            "Vercel 线上测试版可访问基础管理接口，但正式后台仍需要后台登录、数据库、权限、审计和审核队列。",
            "生产化前必须接入数据库、对象存储、真实埋点 SDK、RBAC 权限和内容安全审核。",
        ],
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
